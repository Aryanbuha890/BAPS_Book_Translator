"""
assembler.py

Rebuilds translated chunks into a clean, properly formatted output
document. Deliberately does NOT try to overlay Bengali text into the
original PDF's exact pixel boxes: Bengali script runs wider than
Gujarati/Hindi for equivalent meaning, so forcing it into identical boxes
causes visible overflow/clipping. Instead this preserves structure
(heading levels, paragraph order, chapter boundaries) and lets Bengali
text flow naturally -- this is what actually "looks good" in practice.

Two output formats are supported:
- DOCX (recommended): proper heading styles, native Bengali font support,
  easy to open/print/share.
- TXT: simple fallback, heading markers kept as "== Heading ==".
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


BENGALI_FONT = "Noto Sans Bengali"  # install this font on the machine for correct rendering


def _set_bengali_font(run, size_pt: int = 12):
    run.font.name = BENGALI_FONT
    run.font.size = Pt(size_pt)
    # Word requires the East Asian/complex-script font name set separately
    # for non-Latin scripts to actually render with the chosen font.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:cs"), BENGALI_FONT)
    rfonts.set(qn("w:eastAsia"), BENGALI_FONT)


def assemble_docx(chapters: list[tuple[str, list[str]]], output_path: str) -> str:
    """
    chapters: list of (chapter_title, translated_chunks) where each chunk
    list may contain "" as paragraph-break markers and "## text" as
    heading markers (matching extractor.py's conventions).
    """
    doc = Document()

    for chapter_title, chunks in chapters:
        heading = doc.add_heading(chapter_title, level=1)
        for run in heading.runs:
            _set_bengali_font(run, size_pt=16)

        paragraph = doc.add_paragraph()
        for chunk in chunks:
            if chunk == "":
                paragraph = doc.add_paragraph()
                continue
            if chunk.startswith("## "):
                sub = doc.add_heading(chunk[3:], level=2)
                for run in sub.runs:
                    _set_bengali_font(run, size_pt=14)
                paragraph = doc.add_paragraph()
                continue
            run = paragraph.add_run((chunk + " "))
            _set_bengali_font(run, size_pt=12)

        doc.add_page_break()

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def assemble_txt(chapters: list[tuple[str, list[str]]], output_path: str) -> str:
    """Simple UTF-8 text fallback output."""
    lines = []
    for chapter_title, chunks in chapters:
        lines.append(f"== {chapter_title} ==")
        lines.append("")
        buffer = ""
        for chunk in chunks:
            if chunk == "":
                if buffer:
                    lines.append(buffer)
                    lines.append("")
                buffer = ""
            elif chunk.startswith("## "):
                if buffer:
                    lines.append(buffer)
                    lines.append("")
                    buffer = ""
                lines.append(f"-- {chunk[3:]} --")
                lines.append("")
            else:
                buffer += chunk + " "
        if buffer:
            lines.append(buffer)
        lines.append("\n")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return output_path
